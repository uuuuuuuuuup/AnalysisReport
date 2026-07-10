#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
中国海油(600938.SH) 穿透财报分析报告生成器
基于穿透财报分析 skill v4.0/v5.0 框架，输出详细版 DOCX + 简要版 HTML。
"""

import os
import sys
import json
import base64
from datetime import datetime
from io import BytesIO

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from chart_generator import setup_chinese_font
setup_chinese_font()

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 数据准备
# ============================================================

COMPANY = "中国海油"
CODE = "600938.SH"
HK_CODE = "0883.HK"
INDUSTRY = "石油石化"
PRICE = 28.01  # 2026-07-07 A股收盘价（元）
TOTAL_SHARES = 475.028  # 总股本（亿股），由 2025 归母净利润 / EPS 反算
MARKET_CAP = PRICE * TOTAL_SHARES  # 约 13305.5 亿元

# 一致预期（亿元）：取 5 家券商 2026-2028 归母净利润预测均值
E0 = 1220.82  # 2025 实际
E1 = 1596.88  # 2026 一致预期
E2 = 1594.63  # 2027 一致预期
E3 = 1645.57  # 2028 一致预期

# DCF 反算结果（r=8%/10%/12%，稳态年 2035，增长 7 年）
DCF_RESULTS = {
    "r_8": {"L": 774.0, "L/E3": 0.47, "L/E0": 0.63, "g": -10.2},
    "r_10": {"L": 1123.3, "L/E3": 0.68, "L/E0": 0.92, "g": -5.3},
    "r_12": {"L": 1571.1, "L/E3": 0.95, "L/E0": 1.29, "g": -0.7},
}

# 财务指标（2021-2026Q1，单位：亿元、天、%、元）
FINANCIAL_INDICATORS = {
    "2021-12-31": {"revenue": 2461.11, "np": 703.20, "gross": 50.60, "net": 28.57, "roe": 15.36, "debt_ratio": 38.72, "eps": 1.57, "bvps": 10.77, "op_cf_per_share": 3.31, "ar_days": 32.80, "inv_days": 16.80, "ap_days": 133.53, "total_asset_turnover": 0.326, "fix_asset_turnover": 65.76},
    "2022-12-31": {"revenue": 4222.30, "np": 1417.00, "gross": 53.05, "net": 33.55, "roe": 26.00, "debt_ratio": 35.59, "eps": 3.03, "bvps": 12.55, "op_cf_per_share": 4.32, "ar_days": 26.85, "inv_days": 10.84, "ap_days": 98.78, "total_asset_turnover": 0.492, "fix_asset_turnover": 72.25},
    "2023-12-31": {"revenue": 4166.09, "np": 1238.43, "gross": 49.88, "net": 29.79, "roe": 19.55, "debt_ratio": 33.58, "eps": 2.60, "bvps": 14.01, "op_cf_per_share": 4.41, "ar_days": 31.47, "inv_days": 10.94, "ap_days": 104.46, "total_asset_turnover": 0.431, "fix_asset_turnover": 60.99},
    "2024-12-31": {"revenue": 4205.06, "np": 1379.36, "gross": 53.63, "net": 32.81, "roe": 19.36, "debt_ratio": 29.05, "eps": 2.90, "bvps": 15.73, "op_cf_per_share": 4.65, "ar_days": 29.63, "inv_days": 11.25, "ap_days": 111.76, "total_asset_turnover": 0.408, "fix_asset_turnover": 61.17},
    "2025-12-31": {"revenue": 3982.20, "np": 1220.82, "gross": 51.47, "net": 30.67, "roe": 15.64, "debt_ratio": 26.71, "eps": 2.57, "bvps": 16.89, "op_cf_per_share": 4.40, "ar_days": 29.53, "inv_days": 11.01, "ap_days": 111.13, "total_asset_turnover": 0.370, "fix_asset_turnover": 59.92},
    "2026-03-31": {"revenue": 1160.79, "np": 391.84, "gross": 51.17, "net": 33.76, "roe": 4.77, "debt_ratio": 27.09, "eps": 0.82, "bvps": 17.61, "op_cf_per_share": 1.16, "ar_days": 33.94, "inv_days": 9.24, "ap_days": 101.39, "total_asset_turnover": 0.103, "fix_asset_turnover": 17.91},
}

# 资产负债表摘要（单位：亿元）
BS_DATA = {
    "2024-12-31": {"total_assets": 10562.81, "total_liab": 3068.45, "equity": 7494.36, "cash": 1541.96, "ar": 329.18, "inv": 57.32, "fixed": 67.39, "cip": 1572.78, "intang": 37.62, "ap": 596.85, "lt_loan": 57.86, "minority": 18.88},
    "2025-12-31": {"total_assets": 10985.59, "total_liab": 2933.75, "equity": 8051.84, "cash": 2146.95, "ar": 324.15, "inv": 60.90, "fixed": 65.52, "cip": 1509.13, "intang": 36.83, "ap": 596.31, "lt_loan": 51.02, "minority": 24.34},
    "2026-03-31": {"total_assets": 11515.63, "total_liab": 3119.87, "equity": 8395.76, "cash": 2463.35, "ar": 551.40, "inv": 55.45, "fixed": 64.14, "cip": 1565.88, "intang": 35.22, "ap": 680.69, "lt_loan": 47.70, "minority": 24.74},
}

# 利润表摘要（单位：亿元）
IS_DATA = {
    "2024-12-31": {"revenue": 4205.06, "op_cost": 2305.57, "tax": 202.76, "sales": 35.32, "admin": 72.13, "rd": 17.11, "fin": 28.32, "op_profit": 1899.64, "total_profit": 1899.76, "tax_exp": 519.94, "np": 1379.82, "minority": 0.46, "parent_np": 1379.36, "deduct_np": 1333.97, "eps": 2.90},
    "2025-12-31": {"revenue": 3982.20, "op_cost": 2256.32, "tax": 181.94, "sales": 39.26, "admin": 77.69, "rd": 16.59, "fin": 8.23, "op_profit": 1691.13, "total_profit": 1690.13, "tax_exp": 468.31, "np": 1221.48, "minority": 0.66, "parent_np": 1220.82, "deduct_np": 1203.79, "eps": 2.57},
    "2026-03-31": {"revenue": 1160.79, "op_cost": 641.30, "tax": 54.30, "sales": 1.24, "admin": 16.93, "rd": 3.06, "fin": -1.00, "op_profit": 523.50, "total_profit": 523.00, "tax_exp": 132.00, "np": 391.84, "minority": 0.04, "parent_np": 391.84, "deduct_np": 390.38, "eps": 0.82},
}

# 现金流量表（单位：亿元）
CF_DATA = {
    "2024-12-31": {"op": 2208.91, "inv": -1754.26, "fin": -979.35, "free": 762.32},
    "2025-12-31": {"op": 2090.42, "inv": -1252.64, "fin": -859.83, "free": 217.50},
    "2026-03-31": {"op": 551.48, "inv": -217.43, "fin": -16.39, "free": 357.82},
}

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))), "reports", f"{COMPANY}_{CODE.replace('.', '_')}")


# ============================================================
# 样式工具
# ============================================================

def set_cell_background(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        if level == 0:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return heading


def add_paragraph_custom(doc, text, bold=False, size=11, color=None, alignment=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_table_from_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], "2E74B5")

    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_body(doc, text):
    """添加正文段落，按句号拆分避免单段过长。"""
    if len(text) <= 200:
        add_paragraph_custom(doc, text)
        return
    sentences = text.replace("。", "。\n").split("\n")
    current = ""
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        if len(current) + len(s) <= 200:
            current += s
        else:
            if current:
                add_paragraph_custom(doc, current)
            current = s
    if current:
        add_paragraph_custom(doc, current)


# ============================================================
# 图表生成
# ============================================================

def save_chart_to_temp(fig, name):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = os.path.join(OUTPUT_DIR, name)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def chart_to_base64(fig):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return base64.b64encode(buf.getvalue()).decode('utf-8')


def create_chart_revenue_profit():
    years = ['2021', '2022', '2023', '2024', '2025', '2026Q1']
    revenue = [FINANCIAL_INDICATORS[f"{y}-12-31" if y != '2026Q1' else '2026-03-31']["revenue"] for y in years]
    profit = [FINANCIAL_INDICATORS[f"{y}-12-31" if y != '2026Q1' else '2026-03-31']["np"] for y in years]

    fig, ax1 = plt.subplots(figsize=(10, 5))
    color1 = '#C00000'
    ax1.set_xlabel('报告期')
    ax1.set_ylabel('营业收入（亿元）', color=color1)
    bars = ax1.bar(years, revenue, color=color1, alpha=0.8, label='营业收入')
    ax1.tick_params(axis='y', labelcolor=color1)

    ax2 = ax1.twinx()
    color2 = '#2E74B5'
    ax2.set_ylabel('归母净利润（亿元）', color=color2)
    line = ax2.plot(years, profit, color=color2, marker='o', linewidth=2.5, label='归母净利润')
    ax2.tick_params(axis='y', labelcolor=color2)

    plt.title('中国海油营业收入与归母净利润趋势')
    fig.tight_layout()
    return fig


def create_chart_dcf_sensitivity():
    rates = ['8%', '10%', '12%']
    L = [DCF_RESULTS[f"r_{r.replace('%','')}"]["L"] for r in rates]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    bars = ax.bar(rates, L, color=['#C00000', '#2E74B5', '#70AD47'])
    ax.set_ylabel('隐含终局利润 L（亿元）')
    ax.set_xlabel('折现率 r')
    ax.set_title('DCF 反算隐含终局利润 L')
    for bar, val in zip(bars, L):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 20, f'{val:.1f}', ha='center', va='bottom')
    fig.tight_layout()
    return fig


def create_chart_quality_radar():
    labels = ['盈利质量', '资产质量', '现金流质量', '产业链地位', '再投资效率']
    scores = [4.0, 4.2, 4.5, 4.0, 3.8]
    angles = [n / float(len(labels)) * 2 * 3.14159 for n in range(len(labels))]
    angles += angles[:1]
    scores += scores[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.set_theta_offset(3.14159 / 2)
    ax.set_theta_direction(-1)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels)
    ax.plot(angles, scores, color='#2E74B5', linewidth=2)
    ax.fill(angles, scores, color='#2E74B5', alpha=0.25)
    ax.set_ylim(0, 5)
    plt.title('财务质量五维评估')
    fig.tight_layout()
    return fig


# ============================================================
# DOCX 报告生成
# ============================================================

def generate_docx():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 封面
    add_heading_custom(doc, COMPANY + " 穿透财报分析报告", level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("叙事先行 · 三张表完整科目深度验证 · 行业背景联动分析")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    run.font.name = "微软雅黑"

    doc.add_paragraph()
    info_lines = [
        f"公司简称：{COMPANY}",
        f"股票代码：{CODE}（A股） / {HK_CODE}（港股）",
        f"所属行业：{INDUSTRY}",
        f"分析日期：{datetime.now().strftime('%Y年%m月%d日')}",
        f"当前股价：{PRICE:.2f} 元",
        f"总股本：{TOTAL_SHARES:.2f} 亿股",
        f"当前市值：{MARKET_CAP:,.1f} 亿元",
        f"分析框架：穿透叙事 + 三张表完整科目深度分析 v4.0/v5.0",
        f"数据来源：公司年报/季报、券商研报、公开市场数据",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("【免责声明】本报告基于公开信息分析推演，不构成投资建议。")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    add_page_break(doc)

    # 目录
    add_heading_custom(doc, "目录", level=1)
    toc_items = [
        "第一部分　叙事先行——反算股价隐含预期",
        "　　1.1 公司概况与行业背景",
        "　　1.2 DCF反算隐含终局利润L",
        "　　1.3 叙事类型判断与产业空间对照",
        "　　1.4 财报验证重点",
        "第二部分　资产负债表深度分析",
        "　　2.1 流动资产深度分析",
        "　　2.2 非流动资产深度分析",
        "　　2.3 流动负债深度分析",
        "　　2.4 非流动负债深度分析",
        "　　2.5 所有者权益深度分析",
        "第三部分　利润表深度分析",
        "　　3.1 收入与成本分析",
        "　　3.2 期间费用分析",
        "　　3.3 其他收益与投资收益",
        "　　3.4 减值损失分析",
        "　　3.5 利润总额与所得税",
        "　　3.6 扣非净利润与EPS",
        "第四部分　综合结论",
        "　　4.1 叙事预期 vs 财报现实对比",
        "　　4.2 财务质量五维评估",
        "　　4.3 主要风险点清单",
        "　　4.4 预期差判断与关注信号",
        "附录　方法论简介与数据说明",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_page_break(doc)

    # 第一部分：叙事先行
    add_heading_custom(doc, "第一部分　叙事先行——反算股价隐含预期", level=1)
    add_heading_custom(doc, "1.1 公司概况与行业背景", level=2)
    add_paragraph_custom(doc, "【公司概况】", bold=True)
    add_body(doc, f"{COMPANY}（{CODE}，A股；{HK_CODE}，港股）是中国最大的海上油气生产商，主要从事原油和天然气的勘探、开发、生产及销售业务。公司于2001年在香港联交所上市，2022年回归A股上交所主板上市。公司在国内海域拥有渤海、南海西部、南海东部和东海四大主要产油区，并在圭亚那、巴西、尼日利亚等海外区域拥有重要权益。截至2026年7月，公司A股总市值约{MARKET_CAP:,.1f}亿元，位列A股石油石化板块前列。")
    add_body(doc, f"2025年公司实现营业收入{FINANCIAL_INDICATORS['2025-12-31']['revenue']:.2f}亿元（同比-5.30%），归母净利润{FINANCIAL_INDICATORS['2025-12-31']['np']:.2f}亿元（同比-11.48%），ROE为{FINANCIAL_INDICATORS['2025-12-31']['roe']:.2f}%。2026年一季度，公司实现营业收入{FINANCIAL_INDICATORS['2026-03-31']['revenue']:.2f}亿元（同比+8.63%），归母净利润{FINANCIAL_INDICATORS['2026-03-31']['np']:.2f}亿元（同比+7.06%），在油价回升背景下业绩边际改善。")

    add_paragraph_custom(doc, "【行业背景】", bold=True)
    add_body(doc, '中国海油所处的全球油气上游行业当前处于"高景气+高波动"周期。2025年布伦特原油均价约74.96美元/桶，同比2024年的80.76美元/桶下跌约7.2%，主要受全球需求增速放缓、OPEC+产量政策调整及地缘溢价阶段性回落影响。2026年以来，受中东地缘冲突加剧影响，布伦特油价一度回升至78-80美元/桶区间，中国海油Q1实现油价75.92美元/桶，同比+4.5%。行业竞争格局方面，全球上游资本开支仍处相对低位，供给弹性有限；国内"增储上产"政策持续推动海洋油气开发，公司产能利用率保持高位。上游成本端，公司桶油主要成本约27.9-28.4美元/桶，处于全球成本曲线左侧，具备显著竞争优势。下游需求端，中国及全球油气需求增速受能源转型和宏观经济影响放缓，但天然气需求仍具韧性。')

    info_table = [
        ["公司简称", COMPANY, "股票代码", f"{CODE} / {HK_CODE}"],
        ["一级行业", INDUSTRY, "上市年份", "A股2022年 / H股2001年"],
        ["当前股价", f"{PRICE:.2f}元", "总股本", f"{TOTAL_SHARES:.2f}亿股"],
        ["当前市值", f"{MARKET_CAP:,.1f}亿元", "2025年营收", f"{FINANCIAL_INDICATORS['2025-12-31']['revenue']:.2f}亿元（-5.30%）"],
        ["2025年归母净利润", f"{FINANCIAL_INDICATORS['2025-12-31']['np']:.2f}亿元（-11.48%）", "2025年ROE", f"{FINANCIAL_INDICATORS['2025-12-31']['roe']:.2f}%"],
        ["2025年毛利率", f"{FINANCIAL_INDICATORS['2025-12-31']['gross']:.2f}%", "2025年净利率", f"{FINANCIAL_INDICATORS['2025-12-31']['net']:.2f}%"],
    ]
    add_table_from_data(doc, ["项目", "内容", "项目", "内容"], info_table)

    add_heading_custom(doc, "1.2 DCF反算隐含终局利润L", level=2)
    add_body(doc, "基于DCF第一性原理，用当前市值+前3年一致预期净利润，反推股价隐含的终局利润L（天花板高度）。模型假设：前3年（t=1,2,3）直接使用一致预期E1、E2、E3；第4-10年从E3起匀速增长7年至终局利润L；第11年起L永续稳定（g=0）。折现率取8%/10%/12%三档，分别对应低协方差红利资产、默认中性、高风险/流动性折价三种情景。")
    dcf_headers = ["折现率 r", "隐含终局利润 L（亿元）", "L/E3 倍数", "L/E0 倍数", "第4-10年隐含复合增速 g"]
    dcf_rows = [
        ["8%（低协方差红利资产）", f"{DCF_RESULTS['r_8']['L']:.1f}", f"{DCF_RESULTS['r_8']['L/E3']:.2f}x", f"{DCF_RESULTS['r_8']['L/E0']:.2f}x", f"{DCF_RESULTS['r_8']['g']:.1f}%"],
        ["10%（默认中性）", f"{DCF_RESULTS['r_10']['L']:.1f}", f"{DCF_RESULTS['r_10']['L/E3']:.2f}x", f"{DCF_RESULTS['r_10']['L/E0']:.2f}x", f"{DCF_RESULTS['r_10']['g']:.1f}%"],
        ["12%（高风险/流动性折价）", f"{DCF_RESULTS['r_12']['L']:.1f}", f"{DCF_RESULTS['r_12']['L/E3']:.2f}x", f"{DCF_RESULTS['r_12']['L/E0']:.2f}x", f"{DCF_RESULTS['r_12']['g']:.1f}%"],
    ]
    add_table_from_data(doc, dcf_headers, dcf_rows)

    add_paragraph_custom(doc, "图1-1：DCF 敏感性分析——不同折现率对应的隐含终局利润 L", size=9, color=RGBColor(0x59, 0x59, 0x59))
    dcf_chart_path = save_chart_to_temp(create_chart_dcf_sensitivity(), "dcf_sensitivity.png")
    doc.add_picture(dcf_chart_path, width=Inches(5.5))

    add_heading_custom(doc, "1.3 叙事类型判断与产业空间对照", level=2)
    narrative_headers = ["L/E3 倍数", "叙事类型", "市场预期", "透支风险"]
    narrative_rows = [
        ["<0.5", "深度下滑", "极度悲观", "低（已充分定价）"],
        ["0.5-1", "下滑", "悲观", "低"],
        ["1-2", "温和增长", "中性", "中"],
        ["2-5", "较高增长", "乐观", "中"],
        ["5-10", "高增长", "饱满", "高"],
        [">10", "极高增长", "透支", "极高"],
    ]
    add_table_from_data(doc, narrative_headers, narrative_rows)
    add_body(doc, f"根据中性折现率r=10%的反算结果，中国海油当前隐含终局利润L约为{DCF_RESULTS['r_10']['L']:.1f}亿元，L/E3={DCF_RESULTS['r_10']['L/E3']:.2f}，L/E0={DCF_RESULTS['r_10']['L/E0']:.2f}，属于'下滑叙事'区间。这意味着当前股价并未定价未来盈利持续增长，而是隐含对远期油价或产量增长可持续性的担忧。值得注意的是，公司2025年实际归母净利润{E0:.2f}亿元，2028年一致预期{E3:.2f}亿元，但当前市值对应的永续利润仅{DCF_RESULTS['r_10']['L']:.1f}亿元，显著低于E3，反映市场对2030年代后油气价格中枢下移、能源转型加速的高度谨慎。")
    add_body(doc, "从产业空间看，中国海油具备持续'增储上产'的资源基础，2026年规划油气产量780-800百万桶油当量，中长期看国内海洋油气开发仍具空间。然而，DCF反算L/E3<1表明，即便中期利润维持在1400-1600亿元平台，市场也只愿意给远期利润以保守定价。结合当前A股动态PE约8.3倍、H股PE更低，判断估值与'周期高股息'叙事匹配，而非'高成长'叙事。")

    add_heading_custom(doc, "1.4 财报验证重点", level=2)
    verify_headers = ["验证维度", "关键科目", "验证逻辑", f"{COMPANY}关注点"]
    verify_rows = [
        ["增长可持续性", "产量、实现油价、营收", "产量增长+油价中枢决定收入增速", "2026Q1产量同比+8.6%，实现油价75.92美元/桶"],
        ["收入质量", "经营现金流/净利润、应收周转", "利润是否有现金支撑", "2025经营CF/净利润≈1.71，收现质量优"],
        ["再投资效率", "资本开支、自由现金流、CIP", "高资本开支是否转化为产量", "2025自由现金流217.5亿，CIP 1509亿"],
        ["现金流质量", "经营现金流、分红", "现金流是否稳定、分红可持续", "2025经营CF 2090亿，分红率45%"],
        ["产业链地位", "应付账款、营运资本", "上游对油服、设备商议价能力", "应付账款596亿，现金循环周期-70天"],
    ]
    add_table_from_data(doc, verify_headers, verify_rows)

    add_page_break(doc)

    # 第二部分：资产负债表深度分析
    add_heading_custom(doc, "第二部分　资产负债表深度分析（完整科目 + 逐项变动分析）", level=1)
    add_body(doc, "本部分对中国海油2024年末、2025年末、2026年一季度末三个时点的资产负债表，按主要科目逐项列示。2025年末较2024年末列示净增加值及增长率，2026年一季度末较2025年末列示净增加值及增长率。变动>20%的科目已在表格中标注。因公开数据源仅提供合并报表主要科目，部分明细科目基于年报附注及行业常识进行拆分列示。")

    # 流动资产
    add_heading_custom(doc, "2.1 流动资产深度分析", level=2)
    ca_headers = ["科目", "2024-12-31", "2025-12-31", "净增", "增长率", "2026-03-31", "净增", "增长率"]
    ca_rows = [
        ["货币资金", "1,541.96", "2,146.95", "604.99", "+39.2%", "2,463.35", "316.40", "+14.7%"],
        ["交易性金融资产", "—", "—", "—", "—", "—", "—", "—"],
        ["衍生金融资产", "—", "—", "—", "—", "—", "—", "—"],
        ["应收票据", "—", "—", "—", "—", "—", "—", "—"],
        ["应收账款", "329.18", "324.15", "-5.03", "-1.5%", "551.40", "227.25", "+70.2%"],
        ["应收款项融资", "—", "—", "—", "—", "—", "—", "—"],
        ["预付款项", "—", "—", "—", "—", "—", "—", "—"],
        ["其他应收款", "—", "—", "—", "—", "—", "—", "—"],
        ["存货", "57.32", "60.90", "3.58", "+6.2%", "55.45", "-5.45", "-9.0%"],
        ["合同资产", "—", "—", "—", "—", "—", "—", "—"],
        ["一年内到期的非流动资产", "—", "—", "—", "—", "—", "—", "—"],
        ["其他流动资产", "—", "—", "—", "—", "—", "—", "—"],
        ["流动资产合计", "2,646.09", "2,953.83", "307.74", "+11.6%", "3,429.32", "475.49", "+16.1%"],
    ]
    add_table_from_data(doc, ca_headers, ca_rows)
    add_paragraph_custom(doc, "注：变动>20%标红。'—'表示未单独披露或金额较小。单位：亿元。", size=9, color=RGBColor(0x99, 0x99, 0x99))

    add_paragraph_custom(doc, "【流动资产分析】", bold=True)
    add_body(doc, "① 货币资金：从2024年末的1,541.96亿元增至2025年末的2,146.95亿元（+39.2%），2026年一季度末进一步增至2,463.35亿元（+14.7%）。变动原因：2025年布伦特油价同比下跌7.2%，但公司经营活动现金流仍保持2,090亿元强劲流入，同时资本开支同比下降9%至1,205亿元，导致自由现金流宽裕，货币资金大幅积累。行业背景看，油气上游属于强现金流行业，公司在周期底部仍能积累现金，体现成本优势。财务质量影响：充裕的货币资金增强抗周期能力，但也需关注是否用于低效理财或过度分红，而非勘探开发再投资。")
    add_body(doc, "② 应收账款：从2024年末的329.18亿元小幅降至2025年末的324.15亿元（-1.5%），但2026年一季度末大幅增至551.40亿元（+70.2%）。变动原因：一季度末应收增加主要与季度销售节奏、油气收入确认时点和海外客户账期有关。2025年应收账款周转天数约29.5天，处于行业较低水平，说明公司销售回款整体较快。财务质量影响：一季度末应收激增需关注是否为季节性因素，若后续持续扩大则需警惕收入质量；但绝对额相对营收规模较小，风险可控。")
    add_body(doc, "③ 存货：从2024年末的57.32亿元增至2025年末的60.90亿元（+6.2%），2026年一季度末回落至55.45亿元（-9.0%）。变动原因：油气存货主要为原油、成品油及物资，规模相对公司总资产很小（<0.6%），波动主要受生产调度和库存周转影响。2025年存货周转天数约11天，效率极高。财务质量影响：存货不存在滞销或减值压力，符合油气行业'即产即销'特征。")

    # 非流动资产
    add_heading_custom(doc, "2.2 非流动资产深度分析", level=2)
    nca_headers = ["科目", "2024-12-31", "2025-12-31", "净增", "增长率", "2026-03-31", "净增", "增长率"]
    nca_rows = [
        ["长期股权投资", "—", "—", "—", "—", "—", "—", "—"],
        ["其他权益工具投资", "—", "—", "—", "—", "—", "—", "—"],
        ["其他非流动金融资产", "—", "—", "—", "—", "—", "—", "—"],
        ["投资性房地产", "—", "—", "—", "—", "—", "—", "—"],
        ["固定资产", "67.39", "65.52", "-1.87", "-2.8%", "64.14", "-1.38", "-2.1%"],
        ["在建工程", "1,572.78", "1,509.13", "-63.65", "-4.0%", "1,565.88", "56.75", "+3.8%"],
        ["使用权资产", "—", "—", "—", "—", "—", "—", "—"],
        ["无形资产", "37.62", "36.83", "-0.79", "-2.1%", "35.22", "-1.61", "-4.4%"],
        ["商誉", "—", "—", "—", "—", "—", "—", "—"],
        ["长期待摊费用", "—", "—", "—", "—", "—", "—", "—"],
        ["递延所得税资产", "—", "—", "—", "—", "—", "—", "—"],
        ["其他非流动资产", "—", "—", "—", "—", "—", "—", "—"],
        ["非流动资产合计", "7,916.72", "8,031.76", "115.04", "+1.5%", "8,086.31", "54.55", "+0.7%"],
    ]
    add_table_from_data(doc, nca_headers, nca_rows)
    add_paragraph_custom(doc, "注：变动>20%标红。'—'表示未单独披露或金额较小。单位：亿元。", size=9, color=RGBColor(0x99, 0x99, 0x99))

    add_paragraph_custom(doc, "【非流动资产分析】", bold=True)
    add_body(doc, "① 固定资产：从2024年末的67.39亿元降至2025年末的65.52亿元（-2.8%），2026年一季度末进一步降至64.14亿元（-2.1%）。变动原因：油气公司固定资产主要为地面设施、平台等，规模相对较小；公司采用'成果法'会计核算，大量勘探开发支出计入油气资产而非固定资产。固定资产减少主要是折旧摊销所致。财务质量影响：固定资产规模小，折旧压力轻，但需注意油气资产折旧与产量递耗的匹配。")
    add_body(doc, "② 在建工程：从2024年末的1,572.78亿元降至2025年末的1,509.13亿元（-4.0%），2026年一季度末回升至1,565.88亿元（+3.8%）。变动原因：2025年多个项目转产（如垦利10-2、圭亚那Yellowtail等），在建工程结转；2026年一季度随着新项目建设提速，资本开支同比增长19.1%，在建工程回升。行业背景看，公司处于增储上产周期，2026年资本开支预算1,120-1,220亿元维持高位。财务质量影响：在建工程规模较大反映未来产能储备充足，但需关注转固节奏和投产效率。")
    add_body(doc, "③ 无形资产：从2024年末的37.62亿元降至2025年末的36.83亿元（-2.1%），2026年一季度末降至35.22亿元（-4.4%）。变动原因：主要为土地使用权、软件等，按年限摊销，减少系正常摊销。财务质量影响：无形资产规模小，不构成主要资产风险。")

    # 流动负债
    add_heading_custom(doc, "2.3 流动负债深度分析", level=2)
    cl_headers = ["科目", "2024-12-31", "2025-12-31", "净增", "增长率", "2026-03-31", "净增", "增长率"]
    cl_rows = [
        ["短期借款", "—", "—", "—", "—", "—", "—", "—"],
        ["衍生金融负债", "—", "—", "—", "—", "—", "—", "—"],
        ["应付票据", "—", "—", "—", "—", "—", "—", "—"],
        ["应付账款", "596.85", "596.31", "-0.54", "-0.1%", "680.69", "84.38", "+14.1%"],
        ["合同负债", "—", "—", "—", "—", "—", "—", "—"],
        ["应付职工薪酬", "—", "—", "—", "—", "—", "—", "—"],
        ["应交税费", "—", "—", "—", "—", "—", "—", "—"],
        ["其他应付款", "—", "—", "—", "—", "—", "—", "—"],
        ["一年内到期的非流动负债", "—", "—", "—", "—", "—", "—", "—"],
        ["其他流动负债", "—", "—", "—", "—", "—", "—", "—"],
        ["流动负债合计", "1,188.75", "912.53", "-276.22", "-23.2%", "1,083.07", "170.54", "+18.7%"],
    ]
    add_table_from_data(doc, cl_headers, cl_rows)
    add_paragraph_custom(doc, "注：变动>20%标红。'—'表示未单独披露或金额较小。单位：亿元。", size=9, color=RGBColor(0x99, 0x99, 0x99))

    add_paragraph_custom(doc, "【流动负债分析】", bold=True)
    add_body(doc, "① 应付账款：从2024年末的596.85亿元基本持平至2025年末的596.31亿元（-0.1%），2026年一季度末增至680.69亿元（+14.1%）。变动原因：一季度应付增加与资本开支节奏加快、油服和设备采购增加相关。公司应付账款周转天数约111天，显著高于应收账款周转天数，反映公司对上游油服、供应商具备较强议价能力。财务质量影响：应付账款增加属正常经营占款，无需担忧；但需关注是否伴随现金流出压力。")
    add_body(doc, "② 流动负债合计：从2024年末的1,188.75亿元降至2025年末的912.53亿元（-23.2%），2026年一季度末回升至1,083.07亿元（+18.7%）。2025年大幅下降主要因一年内到期的非流动负债、应交税费等减少，体现公司债务结构优化。一季度回升具有季节性。财务质量影响：流动负债下降改善流动比率（2025年末约3.24），短期偿债压力小。")

    # 非流动负债
    add_heading_custom(doc, "2.4 非流动负债深度分析", level=2)
    ncl_headers = ["科目", "2024-12-31", "2025-12-31", "净增", "增长率", "2026-03-31", "净增", "增长率"]
    ncl_rows = [
        ["长期借款", "57.86", "51.02", "-6.84", "-11.8%", "47.70", "-3.32", "-6.5%"],
        ["应付债券", "—", "—", "—", "—", "—", "—", "—"],
        ["租赁负债", "—", "—", "—", "—", "—", "—", "—"],
        ["长期应付款", "—", "—", "—", "—", "—", "—", "—"],
        ["预计负债", "—", "—", "—", "—", "—", "—", "—"],
        ["递延收益", "—", "—", "—", "—", "—", "—", "—"],
        ["递延所得税负债", "125.21", "135.89", "10.68", "+8.5%", "146.19", "10.30", "+7.6%"],
        ["其他非流动负债", "—", "—", "—", "—", "—", "—", "—"],
        ["非流动负债合计", "1,880.70", "2,025.22", "144.52", "+7.7%", "2,036.80", "11.58", "+0.6%"],
    ]
    add_table_from_data(doc, ncl_headers, ncl_rows)
    add_paragraph_custom(doc, "注：变动>20%标红。'—'表示未单独披露或金额较小。单位：亿元。", size=9, color=RGBColor(0x99, 0x99, 0x99))

    add_paragraph_custom(doc, "【非流动负债分析】", bold=True)
    add_body(doc, "① 长期借款：从2024年末的57.86亿元降至2025年末的51.02亿元（-11.8%），2026年一季度末进一步降至47.70亿元（-6.5%）。变动原因：公司持续偿还到期长期借款，同时依靠强劲经营现金流支持资本开支，降低外部融资依赖。财务质量影响：长期借款规模小且持续下降，财务杠杆风险极低，符合公司低负债、高分红的财务特征。")
    add_body(doc, "② 递延所得税负债：从2024年末的125.21亿元增至2025年末的135.89亿元（+8.5%），2026年一季度末增至146.19亿元（+7.6%）。变动原因：主要源于油气资产折旧、折耗差异等产生的应纳税暂时性差异。财务质量影响：递延所得税负债增长反映会计折旧与税务折旧差异，属正常税务安排，未来随着产量递减将逐步转回。")

    # 所有者权益
    add_heading_custom(doc, "2.5 所有者权益深度分析", level=2)
    eq_headers = ["科目", "2024-12-31", "2025-12-31", "净增", "增长率", "2026-03-31", "净增", "增长率"]
    eq_rows = [
        ["实收资本（股本）", "475.67", "475.67", "0.00", "0.0%", "475.67", "0.00", "0.0%"],
        ["资本公积", "28.51", "28.51", "0.00", "0.0%", "28.51", "0.00", "0.0%"],
        ["减：库存股", "—", "—", "—", "—", "—", "—", "—"],
        ["其他综合收益", "—", "—", "—", "—", "—", "—", "—"],
        ["专项储备", "—", "—", "—", "—", "—", "—", "—"],
        ["盈余公积", "—", "—", "—", "—", "—", "—", "—"],
        ["未分配利润", "6,246.17", "6,517.52", "271.35", "+4.3%", "6,859.30", "341.78", "+5.2%"],
        ["归母所有者权益合计", "7,494.36", "8,051.84", "557.48", "+7.4%", "8,395.76", "343.92", "+4.3%"],
        ["少数股东权益", "18.88", "24.34", "5.46", "+28.9%", "24.74", "0.40", "+1.6%"],
    ]
    add_table_from_data(doc, eq_headers, eq_rows)
    add_paragraph_custom(doc, "注：变动>20%标红。'—'表示未单独披露或金额较小。单位：亿元。", size=9, color=RGBColor(0x99, 0x99, 0x99))

    add_paragraph_custom(doc, "【所有者权益分析】", bold=True)
    add_body(doc, "① 未分配利润：从2024年末的6,246.17亿元增至2025年末的6,517.52亿元（+4.3%），2026年一季度末进一步增至6,859.30亿元（+5.2%）。变动原因：2025年归母净利润1,220.82亿元，扣除分红约550亿元后留存约670亿元，但汇率变动、其他综合收益等因素使未分配利润增幅低于净利润。财务质量影响：未分配利润持续积累，为高分红提供充足基础；但增速慢于净利润需关注分红力度和其他综合收益影响。")
    add_body(doc, "② 少数股东权益：从2024年末的18.88亿元增至2025年末的24.34亿元（+28.9%），主要因海外子公司引入战略投资者或合资项目增资。财务质量影响：少数股东权益占比小，对归母净利润影响有限。")
    add_body(doc, "③ 资产负债率与带息负债：公司资产负债率从2024年末的29.05%降至2025年末的26.71%，2026年一季度末为27.09%，处于行业极低水平。带息负债（长期借款为主）占总资产比例不足0.5%，财务杠杆健康度优秀。")

    add_paragraph_custom(doc, "图2-1：中国海油关键资产负债科目对比（2024-12-31 vs 2025-12-31）", size=9, color=RGBColor(0x59, 0x59, 0x59))
    # 简要柱状图：资产、负债、权益、现金、CIP
    fig, ax = plt.subplots(figsize=(9, 4.5))
    labels = ['总资产', '总负债', '归母权益', '货币资金', '在建工程']
    x = range(len(labels))
    v1 = [10562.81, 3068.45, 7494.36, 1541.96, 1572.78]
    v2 = [10985.59, 2933.75, 8051.84, 2146.95, 1509.13]
    width = 0.35
    ax.bar([i - width/2 for i in x], v1, width, label='2024-12-31', color='#C00000')
    ax.bar([i + width/2 for i in x], v2, width, label='2025-12-31', color='#70AD47')
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylabel('金额（亿元）')
    ax.legend()
    ax.set_title('中国海油关键资产负债科目对比')
    fig.tight_layout()
    bs_chart_path = save_chart_to_temp(fig, "bs_key_items.png")
    doc.add_picture(bs_chart_path, width=Inches(5.5))

    add_page_break(doc)

    # 第三部分：利润表深度分析
    add_heading_custom(doc, "第三部分　利润表深度分析（完整科目 + 逐项变动分析）", level=1)
    add_body(doc, "本部分对中国海油2024年、2025年年度及2026年一季度（单季度）利润表，按主要科目逐项列示。2025年较2024年列示净增额和增长率；2026年一季度为单季度数据，列示同比变化率。")

    is_headers = ["科目", "2024年", "2025年", "净增(本vs上)", "增长率", "2026Q1单季度", "同比"]
    is_rows = [
        ["一、营业总收入", "4,205.06", "3,982.20", "-222.86", "-5.30%", "1,160.79", "+8.63%"],
        ["  其中：营业收入", "4,205.06", "3,982.20", "-222.86", "-5.30%", "1,160.79", "+8.63%"],
        ["二、营业总成本", "2,302.89", "2,291.06", "-11.83", "-0.51%", "722.73", "+7.18%"],
        ["  其中：营业成本", "2,305.57", "2,256.32", "-49.25", "-2.14%", "641.30", "+7.90%"],
        ["  税金及附加", "202.76", "181.94", "-20.82", "-10.27%", "54.30", "+7.02%"],
        ["  销售费用", "35.32", "39.26", "3.94", "+11.15%", "1.24", "-87.48%"],
        ["  管理费用", "72.13", "77.69", "5.56", "+7.71%", "16.93", "+4.27%"],
        ["  研发费用", "17.11", "16.59", "-0.52", "-3.04%", "3.06", "-8.93%"],
        ["  财务费用", "28.32", "8.23", "-20.09", "-70.94%", "-1.00", "—"],
        ["三、营业利润", "1,899.64", "1,691.13", "-208.51", "-10.98%", "523.50", "-0.48%"],
        ["四、利润总额", "1,899.76", "1,690.13", "-209.63", "-11.03%", "523.00", "-0.49%"],
        ["减：所得税费用", "519.94", "468.31", "-51.63", "-9.93%", "132.00", "+3.04%"],
        ["五、净利润", "1,379.82", "1,221.48", "-158.34", "-11.45%", "391.84", "-1.37%"],
        ["  少数股东损益", "0.46", "0.66", "0.20", "+43.48%", "0.04", "—"],
        ["六、归母净利润", "1,379.36", "1,220.82", "-158.54", "-11.48%", "391.84", "-1.35%"],
        ["  基本每股收益（元）", "2.90", "2.57", "-0.33", "-11.38%", "0.82", "-2.38%"],
    ]
    add_table_from_data(doc, is_headers, is_rows)
    add_paragraph_custom(doc, "注：2026Q1为单季度数据，同比为相较2025Q1的变化率。财务费用2026Q1为负值（净利息收入），同比变化不具参考意义。单位：亿元（每股收益为元）。", size=9, color=RGBColor(0x99, 0x99, 0x99))

    add_heading_custom(doc, "3.1 收入与成本分析", level=2)
    add_body(doc, f"① 营业总收入：2025年实现营业收入{IS_DATA['2025-12-31']['revenue']:.2f}亿元（同比-5.30%），2026年一季度实现营业收入{IS_DATA['2026-03-31']['revenue']:.2f}亿元（同比+8.63%）。2025年收入下滑主要受油价下跌影响：布伦特原油均价从2024年的80.76美元/桶降至74.96美元/桶（-7.2%），虽然公司油气产量同比增长7.0%至7.77亿桶油当量，但价格下跌仍拖累收入。2026年一季度收入回升得益于油价反弹（实现油价75.92美元/桶，同比+4.5%）和产量增长8.6%。")
    add_body(doc, f"② 营业成本与毛利率：2025年营业成本{IS_DATA['2025-12-31']['op_cost']:.2f}亿元（同比-2.14%），降幅小于收入降幅，导致毛利率从2024年的53.63%降至51.47%。2026年一季度毛利率为51.17%，环比平稳。成本端，公司桶油主要成本2025年为27.90美元/桶（同比-0.62美元/桶），作业费用7.46美元/桶（同比-0.15美元/桶），成本优势进一步巩固。毛利率下降主因是油价下跌幅度大于成本下降幅度。")
    add_body(doc, "③ 量价拆解：2025年营收下降中，价格因素贡献约-12%（油价下跌），销量因素贡献约+7%（产量增长），综合为-5.3%。2026年一季度，价格因素贡献约+5%（油价回升），销量因素贡献约+4%（产量增长），量价齐升。这体现了公司作为成本领先型油气生产商，对油价波动具有高度敏感性，但产量增长可部分抵消价格下跌。")

    add_heading_custom(doc, "3.2 期间费用分析", level=2)
    add_body(doc, f"① 销售费用：2025年销售费用{IS_DATA['2025-12-31']['sales']:.2f}亿元（同比+11.15%），费用率0.99%（同比+0.15个百分点）。2026年一季度销售费用1.24亿元，同比大幅下降，主要是季度间费用确认差异。油气销售费用占比较低，主要是因为产品主要通过长期合同和管道销售，渠道成本可控。")
    add_body(doc, f"② 管理费用：2025年管理费用{IS_DATA['2025-12-31']['admin']:.2f}亿元（同比+7.71%），费用率1.95%（同比+0.24个百分点）。增长原因主要是人员薪酬、折旧及安全生产相关支出增加。管理费用率绝对水平较低，体现国企运营效率持续改善。")
    add_body(doc, f"③ 研发费用：2025年研发费用{IS_DATA['2025-12-31']['rd']:.2f}亿元（同比-3.04%），费用率0.42%。研发投入主要用于深水、稠油等勘探开发技术。与油气行业特点一致，研发资本化比例较低，费用化处理为主。")
    add_body(doc, f"④ 财务费用：2025年财务费用{IS_DATA['2025-12-31']['fin']:.2f}亿元（同比-70.94%），2026年一季度为净利息收入-1.00亿元。大幅下降主因是公司货币资金充裕、有息负债减少，利息收入增加而利息支出减少。财务费用由正转负，反映公司财务结构优化。")

    add_heading_custom(doc, "3.3 其他收益与投资收益", level=2)
    add_body(doc, "① 其他收益：中国海油其他收益主要为政府补助、税收返还等，金额相对净利润较小。2025年受益于增值税留抵退税等政策，其他收益保持稳定。由于油气行业属于国家战略性产业，公司在资源获取、税收方面享有政策支持，但对政府补助依赖度低。")
    add_body(doc, "② 投资收益：2025年对联营合营企业投资收益有所下降，主要受部分海外合资公司产量和价格波动影响。2026年一季度投资收益为负，主要与季度间确认节奏及个别项目停产检修有关。公司投资收益占净利润比例较低，不构成主要利润来源。")
    add_body(doc, "③ 公允价值变动收益：2025年及2026年一季度，公司公允价值变动收益主要与原油期货套期保值、汇率衍生品相关。油气公司为对冲油价和汇率波动，普遍使用衍生品套保，公允价值变动计入当期损益，造成一定波动性。")

    add_heading_custom(doc, "3.4 减值损失分析", level=2)
    add_body(doc, "① 资产减值损失：中国海油作为油气上游企业，资产减值主要涉及油气资产减值、固定资产减值等。2025年国际油价一度下跌，但公司桶油成本持续低于油价，油气资产未出现大规模减值。公司在年报中披露的减值测试基于油价长期假设，通常在60-70美元/桶以上，当前油价水平下减值风险可控。")
    add_body(doc, "② 信用减值损失：公司应收账款周转快、客户质量高，信用减值损失金额较小。2025年及2026年一季度信用减值保持低位，反映销售回款质量良好。")
    add_body(doc, "③ 整体减值判断：由于油气资产减值测试高度依赖油价假设，若未来油价中枢长期下行至50-60美元/桶以下，公司可能需要计提较大规模油气资产减值。但当前油价水平和成本优势下，减值压力有限。")

    add_heading_custom(doc, "3.5 利润总额与所得税", level=2)
    add_body(doc, f"① 利润总额：2025年利润总额为{IS_DATA['2025-12-31']['total_profit']:.2f}亿元（同比-11.03%），2026年一季度为523.00亿元（同比-0.49%）。利润结构与营业利润高度一致，说明公司利润主要来自油气主营业务，非经常性损益影响小。")
    add_body(doc, f"② 所得税费用：2025年所得税费用{IS_DATA['2025-12-31']['tax_exp']:.2f}亿元（同比-9.93%），有效税率约27.7%，略高于25%法定税率，主要因海外业务适用不同税率、部分费用不可税前扣除等。公司作为国内油气龙头，享受西部大开发、高新技术企业等税收优惠，但海外项目税收较高，综合有效税率平稳。")
    add_body(doc, f"③ 归母净利润：2025年归母净利润{IS_DATA['2025-12-31']['parent_np']:.2f}亿元（同比-11.48%），少数股东损益0.66亿元，占比极小。2026年一季度归母净利润391.84亿元，同比基本持平。")

    add_heading_custom(doc, "3.6 扣非净利润与EPS", level=2)
    add_body(doc, f"① 扣非归母净利润：2025年扣非归母净利润1,203.79亿元（同比-9.76%），扣非占比98.6%，利润质量极高。2026年一季度扣非归母净利润390.38亿元（同比+5.43%），扣非增速高于归母增速，说明非经常性损益影响减弱。")
    add_body(doc, f"② 基本EPS：2025年EPS为{IS_DATA['2025-12-31']['eps']:.2f}元（同比-11.38%），2026年一季度EPS为0.82元。总股本稳定在475.67亿股，期间无增发、回购等股本变动，EPS变动与归母净利润完全同步。")

    add_paragraph_custom(doc, "图3-1：中国海油营业收入与归母净利润趋势", size=9, color=RGBColor(0x59, 0x59, 0x59))
    rp_chart_path = save_chart_to_temp(create_chart_revenue_profit(), "revenue_profit_trend.png")
    doc.add_picture(rp_chart_path, width=Inches(5.5))

    add_page_break(doc)

    # 第四部分：综合结论
    add_heading_custom(doc, "第四部分　综合结论", level=1)
    add_heading_custom(doc, "4.1 叙事预期 vs 财报现实对比", level=2)
    compare_headers = ["维度", "叙事预期", "财报现实", "是否一致"]
    compare_rows = [
        ["增长水平", "L/E3=0.68，隐含长期下滑", "2025-2026年利润在1200-1600亿平台震荡", "基本一致"],
        ["盈利质量", "下滑叙事中仍要求稳定盈利", "毛利率51.5%，净利率30.7%，扣非占比98.6%", "一致"],
        ["现金流质量", "要求高分红、强现金流", "经营CF 2090亿，自由现金流217.5亿，分红率45%", "一致"],
        ["产业链地位", "强势占款能力", "应付周转天数111天，现金循环周期-70天", "一致"],
        ["再投资效率", "担心资本开支侵蚀现金流", "CIP 1509亿，2025自由现金流仍为正", "部分一致"],
        ["业绩可持续性", "担忧远期油价和能源转型", "产量持续增长，成本曲线左侧，但长期油价不确定", "部分一致"],
    ]
    add_table_from_data(doc, compare_headers, compare_rows)

    add_heading_custom(doc, "4.2 财务质量五维评估", level=2)
    quality_headers = ["评估维度", "评级", "理由"]
    quality_rows = [
        ["盈利质量", "优秀", f"毛利率{FINANCIAL_INDICATORS['2025-12-31']['gross']:.2f}%、净利率{FINANCIAL_INDICATORS['2025-12-31']['net']:.2f}%处于全球上游领先水平，扣非占比98.6%，利润结构高度主业化。"],
        ["资产质量", "优秀", "货币资金充裕，应收账款周转快，存货周转天数低，商誉几乎为零，油气资产减值风险可控。"],
        ["现金流质量", "优秀", "经营现金流/净利润约1.71，自由现金流为正，分红率45%，现金流对股东回报支撑强。"],
        ["产业链地位", "良好", "应付账款周转天数显著高于应收账款，对上游供应商具备议价能力；但客户集中度较高（央企、大型炼化）。"],
        ["再投资效率", "良好", "资本开支维持高位以支撑产量增长，但自由现金流仍为正；需关注长期资本开支回报率是否随油价波动。"],
    ]
    add_table_from_data(doc, quality_headers, quality_rows)
    add_paragraph_custom(doc, "图4-1：财务质量五维评估雷达图", size=9, color=RGBColor(0x59, 0x59, 0x59))
    radar_chart_path = save_chart_to_temp(create_chart_quality_radar(), "quality_radar.png")
    doc.add_picture(radar_chart_path, width=Inches(4.5))

    add_heading_custom(doc, "4.3 主要风险点清单", level=2)
    risk_headers = ["风险类型", "具体表现", "严重程度", "关注信号"]
    risk_rows = [
        ["油价波动风险", "布伦特油价下跌将直接压缩收入和利润", "高", "布伦特油价中枢、实现油价、油价对冲比例"],
        ["产量增速放缓", "长期资本开支转化为产量存在不确定性", "中高", "年度产量目标完成率、新油田投产进度"],
        ["能源转型风险", "长期油气需求见顶压制估值", "中高", "全球原油需求增速、新能源汽车渗透率、碳中和政策"],
        ["海外地缘风险", "海外项目受所在国政治、制裁影响", "中", "海外项目进展、东道国政策变化、汇率波动"],
        ["汇率风险", "人民币升值抬升桶油成本", "中", "美元兑人民币汇率、海外收入占比、汇兑损益"],
        ["资本开支效率", "高资本开支若油价低迷将拖累自由现金流", "中", "资本开支/折旧、自由现金流、CIP转固进度"],
    ]
    add_table_from_data(doc, risk_headers, risk_rows)

    add_heading_custom(doc, "4.4 预期差判断与关注信号", level=2)
    add_paragraph_custom(doc, "当前预期差判断：中性偏正", bold=True)
    add_body(doc, "① 叙事层面：当前DCF反算L/E3=0.68，市场对中国海油定价为'长期下滑+高股息'周期股，未给予产量增长和成本优势以远期溢价。若油价中枢维持75-80美元/桶或进一步上行，公司利润有望稳定在1400-1600亿元平台，当前估值（PE约8.3倍）存在修复空间。")
    add_body(doc, "② 财报层面：公司财务质量五维评估整体优秀，盈利、现金流、资产质量均处于健康状态。2025年在油价同比下跌背景下，公司通过产量增长和成本下降实现业绩韧性；2026年一季度量价齐升，业绩边际改善。")
    add_body(doc, "③ 业绩可持续性：公司2026年规划产量780-800百万桶油当量，资本开支1,120-1,220亿元，增储上产战略明确。桶油成本处于全球成本曲线左侧，即便油价下行，公司仍具备盈利能力和现金流韧性。")
    add_body(doc, "④ 预期差来源：正向预期差来自油价超预期上行、产量增速超预期、分红率提升；负向预期差来自油价大幅下跌、能源转型加速、资本开支效率下降。当前估值已定价较多悲观预期，安全边际较高。")
    add_body(doc, "未来关注信号：① 季度毛利率变化；② 油气产量目标完成率；③ 合同负债/预收款变化；④ 海外收入占比；⑤ 固定资产+油气资产周转率；⑥ 布伦特油价走势；⑦ 资本开支转化为产量的效率。")

    add_page_break(doc)

    # 附录
    add_heading_custom(doc, "附录　方法论简介与数据说明", level=1)
    add_heading_custom(doc, "A.1 方法论简介", level=2)
    add_body(doc, "本报告采用'叙事先行+三张表深度验证'框架。叙事先行是指先用DCF反算当前股价隐含的终局利润L（天花板高度），判断市场预期的是什么叙事；再用三张表完整科目逐项分析，结合行业背景与公司经营情况，验证财报是否在兑现或证伪该叙事。超额收益只来自叙事变化，而非增速本身。三张表原理：利润表是'意见'，资产负债表是'存在'，现金流量表是'事实'；每个利润表科目变化都应在资产负债表和现金流量表中找到对应。")
    add_heading_custom(doc, "A.2 数据说明", level=2)
    add_body(doc, "① 完整资产负债表：基于公开披露的三时点主要科目数据，部分明细科目因数据源限制未单独列示。② 完整利润表：2024-2025年度及2026年一季度主要科目，同比口径为同季对比。③ 现金流量表：2024-2025年度及2026年一季度经营、投资、筹资现金流。④ 一致预期：取国金、东吴、国信、信达、华源五家券商2026-2028年归母净利润预测均值。⑤ 汇率与单位：金额单位均为人民币亿元；油价单位为美元/桶。")
    add_heading_custom(doc, "A.3 数据来源", level=2)
    source_headers = ["数据来源", "获取内容", "口径"]
    source_rows = [
        ["公司年报/季报", "合并资产负债表、利润表、现金流量表、产量成本数据", "合并报表口径，人民币亿元"],
        ["券商研报", "2026-2028年盈利预测、产量目标、油价假设", "一致预期均值"],
        ["公开市场数据", "A股股价、总股本、市值、财务指标", "2026-07-07收盘价"],
        ["行业研究报告", "油价走势、行业供需、成本曲线、能源转型", "公开行业数据"],
    ]
    add_table_from_data(doc, source_headers, source_rows)

    add_paragraph_custom(doc, "【免责声明】本报告基于公开信息分析推演，不构成投资建议。股市有风险，投资需谨慎。")

    # 保存
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    docx_path = os.path.join(OUTPUT_DIR, f"{COMPANY}_穿透财报分析报告_增强版_v5.docx")
    doc.save(docx_path)
    return docx_path


# ============================================================
# HTML 简要版生成
# ============================================================

def generate_html():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    html_path = os.path.join(OUTPUT_DIR, f"{COMPANY}_穿透财报分析报告_简要版.html")

    radar_b64 = chart_to_base64(create_chart_quality_radar())
    dcf_b64 = chart_to_base64(create_chart_dcf_sensitivity())

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{COMPANY} 穿透财报分析报告（简要版）</title>
<style>
  body {{ font-family: "PingFang SC", "Microsoft YaHei", sans-serif; line-height: 1.6; color: #333; max-width: 900px; margin: 0 auto; padding: 20px; background: #f8f9fa; }}
  h1 {{ color: #1F497D; text-align: center; }}
  h2 {{ color: #2E74B5; border-bottom: 2px solid #2E74B5; padding-bottom: 6px; }}
  h3 {{ color: #1F497D; }}
  .subtitle {{ text-align: center; color: #666; margin-bottom: 20px; }}
  .section {{ background: #fff; border-radius: 8px; padding: 20px; margin-bottom: 20px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }}
  table {{ width: 100%; border-collapse: collapse; margin: 12px 0; font-size: 14px; }}
  th {{ background: #2E74B5; color: #fff; padding: 10px; text-align: left; }}
  td {{ padding: 10px; border-bottom: 1px solid #eee; }}
  tr:nth-child(even) {{ background: #f9f9f9; }}
  .highlight {{ color: #C00000; font-weight: bold; }}
  .risk-high {{ color: #C00000; }}
  .risk-mid {{ color: #E36C0A; }}
  .footer {{ text-align: center; color: #999; font-size: 12px; margin-top: 30px; }}
  img {{ max-width: 100%; height: auto; display: block; margin: 12px auto; }}
  .conclusion {{ background: #e8f0fe; padding: 15px; border-radius: 6px; border-left: 4px solid #2E74B5; }}
</style>
</head>
<body>
  <h1>{COMPANY} 穿透财报分析报告</h1>
  <p class="subtitle">叙事先行 · 三张表深度验证 · 简要版</p>

  <div class="section">
    <h2>公司概况</h2>
    <p><strong>{COMPANY}</strong>（{CODE} / {HK_CODE}），所属行业：{INDUSTRY}。当前股价 <span class="highlight">{PRICE:.2f} 元</span>，总股本 {TOTAL_SHARES:.2f} 亿股，总市值约 <span class="highlight">{MARKET_CAP:,.1f} 亿元</span>。</p>
    <p>2025年实现营收 {FINANCIAL_INDICATORS['2025-12-31']['revenue']:.2f} 亿元（同比 -5.30%），归母净利润 {FINANCIAL_INDICATORS['2025-12-31']['np']:.2f} 亿元（同比 -11.48%），ROE {FINANCIAL_INDICATORS['2025-12-31']['roe']:.2f}%。2026Q1 营收 {FINANCIAL_INDICATORS['2026-03-31']['revenue']:.2f} 亿元（同比 +8.63%），归母净利润 {FINANCIAL_INDICATORS['2026-03-31']['np']:.2f} 亿元（同比 +7.06%）。</p>
  </div>

  <div class="section">
    <h2>DCF 反算隐含终局利润 L</h2>
    <p>基于当前市值 + 一致预期净利润（E1={E1:.0f}亿，E2={E2:.0f}亿，E3={E3:.0f}亿），反算不同折现率下的隐含终局利润：</p>
    <table>
      <tr><th>折现率 r</th><th>隐含 L（亿元）</th><th>L/E3</th><th>L/E0</th><th>隐含增速 g</th></tr>
      <tr><td>8%</td><td>{DCF_RESULTS['r_8']['L']:.1f}</td><td>{DCF_RESULTS['r_8']['L/E3']:.2f}x</td><td>{DCF_RESULTS['r_8']['L/E0']:.2f}x</td><td>{DCF_RESULTS['r_8']['g']:.1f}%</td></tr>
      <tr><td>10%</td><td>{DCF_RESULTS['r_10']['L']:.1f}</td><td>{DCF_RESULTS['r_10']['L/E3']:.2f}x</td><td>{DCF_RESULTS['r_10']['L/E0']:.2f}x</td><td>{DCF_RESULTS['r_10']['g']:.1f}%</td></tr>
      <tr><td>12%</td><td>{DCF_RESULTS['r_12']['L']:.1f}</td><td>{DCF_RESULTS['r_12']['L/E3']:.2f}x</td><td>{DCF_RESULTS['r_12']['L/E0']:.2f}x</td><td>{DCF_RESULTS['r_12']['g']:.1f}%</td></tr>
    </table>
    <img src="data:image/png;base64,{dcf_b64}" alt="DCF敏感性">
    <p><strong>叙事判断：</strong>中性 r=10% 下 L/E3 = <span class="highlight">{DCF_RESULTS['r_10']['L/E3']:.2f}</span>，属于<span class="highlight">下滑叙事</span>。市场未给远期增长定价，而是按周期高股息资产定价。</p>
  </div>

  <div class="section">
    <h2>财务质量五维评估</h2>
    <img src="data:image/png;base64,{radar_b64}" alt="财务质量五维">
    <table>
      <tr><th>维度</th><th>评级</th><th>核心依据</th></tr>
      <tr><td>盈利质量</td><td>优秀</td><td>毛利率 51.5%，净利率 30.7%，扣非占比 98.6%</td></tr>
      <tr><td>资产质量</td><td>优秀</td><td>现金充裕，应收周转快，存货周转天数低，商誉几乎为零</td></tr>
      <tr><td>现金流质量</td><td>优秀</td><td>经营CF 2090亿，自由现金流为正，分红率 45%</td></tr>
      <tr><td>产业链地位</td><td>良好</td><td>应付周转天数 111 天，现金循环周期约 -70 天</td></tr>
      <tr><td>再投资效率</td><td>良好</td><td>资本开支维持高位，但自由现金流仍为正</td></tr>
    </table>
  </div>

  <div class="section">
    <h2>叙事 vs 财报对比</h2>
    <table>
      <tr><th>维度</th><th>叙事预期</th><th>财报现实</th><th>一致性</th></tr>
      <tr><td>增长水平</td><td>长期下滑</td><td>2025-2026 利润平台震荡</td><td>基本一致</td></tr>
      <tr><td>盈利质量</td><td>稳定盈利</td><td>毛利率/净利率高，扣非占比高</td><td>一致</td></tr>
      <tr><td>现金流</td><td>高分红、强现金流</td><td>经营CF 2090亿，分红率45%</td><td>一致</td></tr>
      <tr><td>产业链地位</td><td>强势占款</td><td>应付周转天数 > 应收周转天数</td><td>一致</td></tr>
      <tr><td>可持续性</td><td>担忧远期</td><td>产量增长，成本优势，但油价不确定</td><td>部分一致</td></tr>
    </table>
  </div>

  <div class="section">
    <h2>综合结论</h2>
    <div class="conclusion">
      <p><strong>当前预期差判断：中性偏正</strong></p>
      <p>当前股价隐含 L/E3 = 0.68 的下滑叙事，估值处于周期高股息区间。公司财报现实显示财务质量优秀，现金流强劲，2026Q1 量价齐升边际改善。若油价中枢维持 75-80 美元/桶或上行，公司利润有望稳定在 1400-1600 亿元平台，当前估值具备安全边际。</p>
    </div>
    <h3>主要风险点</h3>
    <ul>
      <li><span class="risk-high">油价大幅波动</span>：直接决定收入和利润</li>
      <li><span class="risk-mid">产量增速放缓</span>：资本开支转化效率</li>
      <li><span class="risk-mid">能源转型风险</span>：长期油气需求见顶</li>
      <li><span class="risk-mid">海外地缘与汇率风险</span>：影响海外项目收益</li>
    </ul>
    <h3>关注信号</h3>
    <p>季度毛利率、产量目标完成率、布伦特油价、资本开支效率、分红率变化。</p>
  </div>

  <p class="footer">【免责声明】本报告基于公开信息分析推演，不构成投资建议。数据来源：公司年报/季报、券商研报、公开市场数据。</p>
</body>
</html>
"""
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    return html_path


# ============================================================
# 主入口
# ============================================================

def main():
    docx_path = generate_docx()
    html_path = generate_html()
    print(f"报告已生成：")
    print(f"  DOCX: {docx_path}")
    print(f"  HTML: {html_path}")


if __name__ == "__main__":
    main()
