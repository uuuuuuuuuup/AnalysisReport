from docx import Document

doc_path = "/Users/apple/Documents/分析报告/.claude/reports/中国海油_600938_SH/中国海油_穿透财报分析报告_增强版_v5.docx"
doc = Document(doc_path)
print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")
print(f"总字符数: {sum(len(p.text) for p in doc.paragraphs)}")
