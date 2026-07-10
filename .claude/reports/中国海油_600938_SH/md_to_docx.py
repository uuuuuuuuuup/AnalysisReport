#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert generated Markdown report to DOCX with v4.0 formatting."""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

OUTPUT_DIR = "/Users/apple/Documents/分析报告/.claude/reports/中国海油_600938_SH"
MD_PATH = os.path.join(OUTPUT_DIR, "中国海油_穿透财报分析报告.md")
DOCX_PATH = os.path.join(OUTPUT_DIR, "中国海油_穿透财报分析报告_增强版_v5.docx")


def set_chinese_font(run, font_name="宋体", size=12, bold=False, color=None):
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_heading(text, level=1)
        for run in p.runs:
            set_chinese_font(run, "黑体", 18, bold=True)
    elif level == 2:
        p = doc.add_heading(text, level=2)
        for run in p.runs:
            set_chinese_font(run, "黑体", 14, bold=True)
    else:
        p = doc.add_heading(text, level=3)
        for run in p.runs:
            set_chinese_font(run, "黑体", 12, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5


def add_paragraph(doc, text, bold=False, alignment=None, color=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    set_chinese_font(run, "宋体", 12, bold=bold, color=color)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(24) if not bold else Pt(0)
    return p


def add_table_from_md(doc, lines):
    """Parse markdown table lines and add a Word table."""
    header = lines[0].strip().split("|")[1:-1]
    header = [h.strip() for h in header]
    rows = []
    for line in lines[2:]:
        if not line.strip():
            continue
        cells = line.strip().split("|")[1:-1]
        cells = [c.strip() for c in cells]
        rows.append(cells)
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, "黑体", 10, bold=True)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            row_cells[c_idx].text = cell_text
            for paragraph in row_cells[c_idx].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, "宋体", 9)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()


def add_image(doc, img_path):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()


def convert():
    doc = Document()
    # Set default font for document
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    with open(MD_PATH, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    i = 0
    table_lines = []
    while i < len(lines):
        line = lines[i]
        # Title
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
        elif line.startswith("!"):
            # Image reference
            match = re.match(r"!\[.*?\]\((.*?)\)", line)
            if match:
                img_name = match.group(1)
                img_path = os.path.join(OUTPUT_DIR, img_name)
                add_image(doc, img_path)
        elif line.startswith("| "):
            table_lines.append(line)
            if i + 1 < len(lines) and not lines[i + 1].startswith("| "):
                add_table_from_md(doc, table_lines)
                table_lines = []
        elif line.startswith("**") and line.endswith("**"):
            # Subtitle or bold paragraph
            text = line.strip("*")
            if len(text) < 50:
                add_paragraph(doc, text, bold=True)
            else:
                add_paragraph(doc, text, bold=False)
        elif line.strip() == "":
            pass
        else:
            # Check for inline bold
            add_paragraph(doc, line)
        i += 1
    # Handle last table if any
    if table_lines:
        add_table_from_md(doc, table_lines)

    # Disclaimer on cover page? Already in md.
    doc.save(DOCX_PATH)
    print(f"Word 报告已生成: {DOCX_PATH}")


if __name__ == "__main__":
    convert()
